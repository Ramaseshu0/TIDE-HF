"""
src/chf_titration/rag.py
========================
Local RAG engine for TIDE-HF.

Free · local · no training · no API keys · no internet after first install.

Stack
-----
  ChromaDB             vector store  — ~50 MB on disk
  all-MiniLM-L6-v2    embeddings    — ~80 MB, CPU, downloads once
  Ollama + mistral     generation    — ~4.1 GB, downloads once

Quick start
-----------
  pip install -e ".[rag]"
  ollama pull mistral
  python scripts/setup_rag.py          # indexes rag_docs/ once
  # then open the Streamlit UI and use the Ask AI tab
"""

from __future__ import annotations

import os
import textwrap
from pathlib import Path
from typing import Optional


# ── repo-relative paths ───────────────────────────────────────────────────────
# src/chf_titration/rag.py  →  repo root is four levels up
_HERE     = Path(__file__).parent                        # src/chf_titration/
_SRC      = _HERE.parent                                 # src/
_ROOT     = _SRC.parent                                  # repo root
_DB_DIR   = Path(os.environ.get("TIDE_RAG_DB",   str(_ROOT / "rag_db")))
_DOCS_DIR = Path(os.environ.get("TIDE_RAG_DOCS", str(_ROOT / "rag_docs")))

# ── tunables via env ──────────────────────────────────────────────────────────
_EMBED_MODEL  = os.environ.get("TIDE_EMBED_MODEL",  "all-MiniLM-L6-v2")
_OLLAMA_MODEL = os.environ.get("TIDE_OLLAMA_MODEL", "qwen3.5:4b")
_OLLAMA_URL   = os.environ.get("TIDE_OLLAMA_URL",   "http://localhost:11434")
_GROQ_MODEL   = os.environ.get("TIDE_GROQ_MODEL",   "llama-3.3-70b-versatile")
_COLLECTION   = "tide_hf_v1"

# Generation params — tuned for patient-personalized answers without
# being so creative they invent clinical facts.
_TEMPERATURE = float(os.environ.get("TIDE_LLM_TEMPERATURE", "0.45"))
_TOP_P       = float(os.environ.get("TIDE_LLM_TOP_P",       "0.9"))
_MAX_TOKENS  = int(  os.environ.get("TIDE_LLM_MAX_TOKENS",  "900"))
_RETRIEVE_K  = int(  os.environ.get("TIDE_RETRIEVE_K",      "10"))

# ── system prompts ────────────────────────────────────────────────────────────
# These prompts are deliberately strict: every answer MUST be grounded in
# this specific patient's numbers, and every clinical claim MUST come from
# the retrieved guideline chunks. No generic "talk to your doctor" filler.

_SYSTEM_PATIENT = textwrap.dedent("""\
    You are TIDE-HF, a heart-failure care companion talking with ONE specific
    patient whose data is provided. Speak warmly and in plain English.

    Rules for every answer:
    • Quote at least TWO of this patient's actual numbers (e.g. "your potassium
      of 6.3", "your weight is up 4 lb"). Never give generic answers.
    • Explain any engine flag (global_stop, hold, order_labs) using the
      patient's numbers, not abstract definitions.
    • Translate jargon: ARNi → "Entresto", MRA → "spironolactone (a water-pill
      helper)", SGLT2i → "Jardiance/Farxiga". Define on first use.
    • End with ONE specific next step tied to a number from the data.
    • Never invent facts not in the patient data or guidelines.

    Length: 3 short paragraphs. Use **bold** for key numbers.""")

_SYSTEM_CLINICIAN = textwrap.dedent("""\
    You are TIDE-HF, a clinical decision-support assistant. The data is from
    ONE specific HFrEF patient: engine flags, lab gates, contraindications,
    per-class actions. Cite only the provided guidelines (AHA/ACC/HFSA 2022,
    STRONG-HF, drug monographs).

    Rules for every answer:
    • Anchor every claim to THIS patient's numbers
      (e.g. "K⁺ 6.3 > 6.0 ⇒ global_stop", "bisoprolol 5 mg = 50% of 10 mg target").
    • When you cite a threshold, give the threshold AND the patient value AND
      the delta/ratio.
    • If an engine action looks counter-intuitive, name the rule that produced
      it (lab gate, AE resolver, dose-rung table, contraindication cascade).
    • End with ONE concrete next step — lab, dose, monitor, or referral.

    Length: ≤220 words. Use Markdown bullets for the action list.""")


# ═════════════════════════════════════════════════════════════════════════════
class TideRAG:
    """
    Local RAG engine.

    Usage
    -----
    rag = TideRAG()
    rag.setup("rag_docs/")          # one-time: chunk + embed all PDFs
    ctx = build_engine_context(...)  # from the engine outputs
    answer = rag.ask("Why was my dose reduced?", ctx, audience="patient")
    """

    def __init__(self) -> None:
        self._check_deps()
        import chromadb
        from sentence_transformers import SentenceTransformer

        _DB_DIR.mkdir(parents=True, exist_ok=True)
        self._embed_model = SentenceTransformer(_EMBED_MODEL)
        self._client = chromadb.PersistentClient(path=str(_DB_DIR))
        self._col = self._client.get_or_create_collection(
            name=_COLLECTION,
            metadata={"hnsw:space": "cosine"},
        )

    # ── one-time setup ────────────────────────────────────────────────────────

    def setup(
        self,
        docs_dir: str | Path = _DOCS_DIR,
        force: bool = False,
    ) -> int:
        """
        Chunk and embed every PDF in docs_dir into ChromaDB.

        Expected layout:
            rag_docs/
              clinical/   ← shown to clinicians  (AHA guidelines, monographs)
              patient/    ← shown to patients    (plain-language guides)
              shared/     ← shown to both        (optional)

        Returns the total number of chunks stored.
        """
        self._check_deps(with_pdf=True)
        docs_dir = Path(docs_dir)

        if not docs_dir.exists():
            raise FileNotFoundError(
                f"Docs folder not found: {docs_dir}\n"
                f"Create it with:\n"
                f"  mkdir -p {docs_dir}/clinical {docs_dir}/patient"
            )

        if self._col.count() > 0 and not force:
            print(
                f"[RAG] {self._col.count()} chunks already indexed. "
                "Pass force=True to re-index from scratch."
            )
            return self._col.count()

        if force and self._col.count() > 0:
            self._client.delete_collection(_COLLECTION)
            self._col = self._client.get_or_create_collection(
                name=_COLLECTION, metadata={"hnsw:space": "cosine"}
            )

        total = 0
        for audience in ("clinical", "patient", "shared"):
            sub = docs_dir / audience
            if not sub.exists():
                continue
            files = sorted(
                list(sub.glob("*.pdf")) + list(sub.glob("*.docx")),
                key=lambda p: p.name.lower(),
            )
            for src in files:
                print(f"  [RAG] indexing  {src.name}  ({audience})")
                try:
                    text = (
                        self._pdf_to_text(src) if src.suffix.lower() == ".pdf"
                        else self._docx_to_text(src)
                    )
                except Exception as exc:
                    print(f"    ⚠ skipped (cannot extract text): {exc}")
                    continue
                chunks = self._chunk(text)
                if not chunks:
                    print(f"    ⚠ skipped (no extractable text)")
                    continue
                for i, chunk in enumerate(chunks):
                    self._col.add(
                        ids        =[f"{audience}__{src.stem}__{i}"],
                        embeddings =[self._encode(chunk)],
                        documents  =[chunk],
                        metadatas  =[{"audience": audience, "source": src.name}],
                    )
                    total += 1

        print(f"[RAG] Indexed {total} chunks → {_DB_DIR}")
        return total

    # ── live query ────────────────────────────────────────────────────────────

    def ask(
        self,
        question:       str,
        engine_context: str,
        audience:       str = "patient",  # "patient" | "clinician"
        k:              int | None = None,
    ) -> str:
        """
        Retrieve relevant guideline chunks, build the prompt, call the LLM.

        The retrieval query is the question PLUS the salient signals from the
        engine context (active flags, dose changes, key labs). This pulls in
        chunks that match what's actually happening with this patient, not
        just the literal words of the question.
        """
        if self._col.count() == 0:
            return (
                "The knowledge base is empty.\n\n"
                "Run  python scripts/setup_rag.py  first to index your PDFs."
            )

        k = k or _RETRIEVE_K
        retrieval_query = self._build_retrieval_query(question, engine_context)

        chunks = self._retrieve(retrieval_query, audience, k)
        if not chunks:
            chunks = self._retrieve(retrieval_query, "shared", k)  # fallback

        context_block = (
            "\n\n---\n\n".join(chunks) if chunks else "(no guidelines retrieved)"
        )
        system = _SYSTEM_PATIENT if audience == "patient" else _SYSTEM_CLINICIAN

        # XML-style tags are unambiguous to small instruction-tuned models
        # (qwen3.5:4b, llama-3.1-8b). They prevent the "=== SECTION ===" header
        # from being misread as a placeholder for content the user hasn't sent.
        prompt = (
            f"{system}\n\n"
            f"You are now reviewing this real patient. Their full state is in "
            f"<patient_data>. Use it directly — do NOT say information is "
            f"missing.\n\n"
            f"<patient_data>\n{engine_context}\n</patient_data>\n\n"
            f"<guidelines>\n{context_block}\n</guidelines>\n\n"
            f"<question audience=\"{audience}\">\n{question}\n</question>\n\n"
            f"Write your answer now. Reference at least two specific numbers "
            f"from <patient_data> in the first paragraph. Quote the relevant "
            f"threshold from <guidelines> when you cite a clinical rule. Do "
            f"not output XML tags or section headers in your answer."
        )

        return self._call_ollama(prompt)

    @staticmethod
    def _build_retrieval_query(question: str, engine_context: str) -> str:
        """Combine the user's question with the salient lines from the engine
        context so retrieval matches what's actually going on with the
        patient, not just the words in the question."""
        salient: list[str] = []
        for line in engine_context.splitlines():
            l = line.strip()
            if not l:
                continue
            # Keep the high-signal lines: status, flags, labs, decisions.
            if l.startswith(("STATUS:", "Active adverse effects:", "Labs",
                              "RAAS:", "Patient:")):
                salient.append(l)
            elif l.lstrip().startswith(("ACEi", "ARB", "ARNi", "RAAS",
                                          "beta_blocker", "MRA", "SGLT2i",
                                          "loop")):
                salient.append(l.strip())
        return question + "\n" + "\n".join(salient[:12])

    # ── internal helpers ──────────────────────────────────────────────────────

    def _encode(self, text: str) -> list[float]:
        return self._embed_model.encode(text, normalize_embeddings=True).tolist()

    def _retrieve(self, query: str, audience: str, k: int) -> list[str]:
        where = (
            {"audience": {"$in": [audience, "shared"]}}
            if audience in ("clinical", "patient")
            else {}
        )
        try:
            res = self._col.query(
                query_embeddings=[self._encode(query)],
                n_results=min(k, self._col.count()),
                where=where or None,
            )
            return res["documents"][0] if res["documents"] else []
        except Exception:
            return []

    def _call_ollama(self, prompt: str) -> str:
        """Call Ollama; on any failure (e.g. macOS Tahoe Metal bug), fall back
        to Groq if GROQ_API_KEY is set."""
        try:
            import requests
            resp = requests.post(
                f"{_OLLAMA_URL}/api/generate",
                json={
                    "model":  _OLLAMA_MODEL,
                    "prompt": prompt,
                    "stream": False,
                    # Disable Qwen-style chain-of-thought tokens — we want the
                    # whole budget spent on the actual answer.
                    "think":  False,
                    "options": {
                        "temperature":     _TEMPERATURE,
                        "top_p":           _TOP_P,
                        "num_predict":     _MAX_TOKENS,
                        "repeat_penalty":  1.15,   # discourage boilerplate
                    },
                },
                timeout=180,
            )
            resp.raise_for_status()
            data = resp.json()
            # Prefer .response. If it's empty (older Ollama still emits the
            # thinking field even when think:false), salvage that as a fallback.
            text = (data.get("response") or "").strip()
            if not text:
                text = (data.get("thinking") or "").strip()
            if text:
                return text
            ollama_err: Exception | str = (
                f"empty response (done_reason={data.get('done_reason')})"
            )
        except Exception as exc:
            ollama_err = exc

        # ── Groq fallback (only if user has set GROQ_API_KEY) ────────────
        groq_key = os.environ.get("GROQ_API_KEY", "").strip()
        if groq_key:
            answer = self._call_groq(prompt, groq_key)
            if answer:
                return answer

        return (
            f"⚠  Local LLM is not reachable ({ollama_err}).\n\n"
            f"Two ways to enable a real RAG response:\n"
            f"  • Ollama (local, no key):   ollama serve  &&  ollama pull {_OLLAMA_MODEL}\n"
            f"  • Groq (hosted, free key):  export GROQ_API_KEY=gsk_…\n\n"
            f"Then send your question again."
        )

    @staticmethod
    def _call_groq(prompt: str, api_key: str) -> str | None:
        """Call Groq's hosted Llama-3.3 with the same RAG prompt. Returns
        None on failure so the caller can surface a clear error."""
        try:
            from groq import Groq
        except ImportError:
            return None
        try:
            client = Groq(api_key=api_key)
            resp = client.chat.completions.create(
                model=_GROQ_MODEL,
                temperature=_TEMPERATURE,
                top_p=_TOP_P,
                max_tokens=_MAX_TOKENS,
                messages=[{"role": "user", "content": prompt}],
            )
            return (resp.choices[0].message.content or "").strip() or None
        except Exception:
            return None

    @staticmethod
    def _pdf_to_text(path: Path) -> str:
        from pypdf import PdfReader
        return "\n".join(
            (p.extract_text() or "") for p in PdfReader(str(path)).pages
        )

    @staticmethod
    def _docx_to_text(path: Path) -> str:
        try:
            from docx import Document
        except ImportError as e:
            raise ImportError(
                "python-docx is required to index .docx files. "
                "Install with: pip install -e '.[rag]'"
            ) from e
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)

    @staticmethod
    def _chunk(text: str, size: int = 400, overlap: int = 50) -> list[str]:
        words = text.split()
        out, i = [], 0
        while i < len(words):
            chunk = " ".join(words[i : i + size])
            if len(chunk.strip()) > 40:
                out.append(chunk)
            i += size - overlap
        return out

    @staticmethod
    def _check_deps(with_pdf: bool = False) -> None:
        missing = []
        for pkg in ["chromadb", "sentence_transformers", "requests"]:
            try:
                __import__(pkg)
            except ImportError:
                missing.append(pkg.replace("_", "-"))
        if with_pdf:
            for pkg, pip_name in (("pypdf", "pypdf"), ("docx", "python-docx")):
                try:
                    __import__(pkg)
                except ImportError:
                    missing.append(pip_name)
        if missing:
            raise ImportError(
                f"Missing RAG dependencies: {', '.join(missing)}\n"
                "Install with:  pip install -e '.[rag]'"
            )

    # ── convenience property ──────────────────────────────────────────────────

    @property
    def chunk_count(self) -> int:
        return self._col.count()


# ═════════════════════════════════════════════════════════════════════════════
def build_engine_context(
    result:  dict,
    changes: dict,
    patient: dict,
    flags:   dict,
    labs:    Optional[dict] = None,
) -> str:
    """
    Convert TitrationEngine + apply_strategy outputs into a plain-text
    context block that the RAG prompt can read.

    Parameters are exactly what ui.py already produces — pass them straight in.
    """
    lines: list[str] = []

    # demographics
    lines.append(
        f"Patient: {patient.get('age', '?')}y "
        f"{patient.get('gender', '?')}  |  HFrEF NYHA C/D"
    )

    # global stop
    if result.get("global_stop"):
        lines.append("STATUS: GLOBAL STOP — all GDMT held, urgent evaluation required")

    # AE flags
    active = [
        f.replace("_detected", "").replace("_", " ")
        for f, v in flags.items() if v
    ]
    lines.append(
        "Active adverse effects: " + (", ".join(active) if active else "none")
    )

    # labs
    if labs:
        bl_cr  = (patient.get("baseline") or {}).get("Cr") or 0
        cr_pct = (labs["Cr"] - bl_cr) / bl_cr * 100 if bl_cr else None
        cr_str = f"{labs['Cr']:.2f}"
        if cr_pct is not None:
            cr_str += f" ({cr_pct:+.0f}% vs baseline)"
        lines.append(
            f"Labs — K+:{labs['K']:.1f} mEq/L  Na:{labs['Na']:.0f}  "
            f"Cr:{cr_str} mg/dL  eGFR:{labs['eGFR']:.0f}  HCO3:{labs['HCO3']:.0f}"
        )

    # labs ordered
    if result.get("order_labs"):
        lines.append(
            "Labs ordered: " + ", ".join(sorted(result.get("labs_requested", [])))
        )

    # RAAS
    lines.append(
        f"RAAS: preferred={result.get('preferred_raas') or 'none'}  "
        f"active={result.get('active_raas') or 'none'}"
    )

    # per-drug decisions
    lines.append("Titration decisions:")
    for cls, c in changes.items():
        action = c.get("concrete_action", "")
        cur    = c.get("current") or 0
        nd     = c.get("new_dose") or 0
        tgt    = c.get("target")  or 0
        reason = c.get("reason", "").replace("_", " ")
        pct    = f"  ({int(nd/tgt*100)}% of target)" if tgt and nd else ""
        lines.append(f"  {cls:<14} {action:<22} {cur}mg → {nd}mg{pct}  [{reason}]")

    if result.get("awaiting_labs_next"):
        lines.append(
            "Awaiting labs next cycle: "
            + ", ".join(result["awaiting_labs_next"])
        )

    return "\n".join(lines)
